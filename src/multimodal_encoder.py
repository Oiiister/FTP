"""
多模态编码器模块
支持文本、图像、表格、PDF等多种数据格式的向量化编码
"""
import os
import torch
import pandas as pd
from typing import Union, List, Dict, Any
from PIL import Image
import pdfplumber

# 导入Hugging Face Transformers库
from transformers import (
    AutoModel, AutoProcessor, AutoTokenizer, AutoImageProcessor,
    CLIPModel, CLIPProcessor
)


class MultimodalEncoder:
    """多模态数据编码器"""
    
    def __init__(self, device: str = "cuda" if torch.cuda.is_available() else "cpu"):
        """
        初始化多模态编码器
        
        Args:
            device: 计算设备 (cuda/cpu)
        """
        self.device = device
        self._load_models()
    
    def _load_models(self):
        """加载各种编码器模型"""
        print("正在加载多模态编码器...")
        
        # 1. 图像编码器 - CLIP (通用视觉编码器)
        try:
            self.clip_model = CLIPModel.from_pretrained("openai/clip-vit-base-patch32")
            self.clip_processor = CLIPProcessor.from_pretrained("openai/clip-vit-base-patch32")
            self.clip_model.to(self.device)
            print("✓ CLIP图像编码器加载成功")
        except Exception as e:
            print(f"⚠ CLIP加载失败: {e}")
            self.clip_model = None
            self.clip_processor = None
        
        # 2. 文本编码器 - BGE-M3 (当前SOTA文本嵌入模型)
        try:
            self.text_model = AutoModel.from_pretrained("BAAI/bge-m3", trust_remote_code=True)
            self.text_tokenizer = AutoTokenizer.from_pretrained("BAAI/bge-m3", trust_remote_code=True)
            self.text_model.to(self.device)
            print("✓ BGE-M3文本编码器加载成功")
        except Exception as e:
            print(f"⚠ BGE-M3加载失败: {e}")
            # 备用文本编码器
            try:
                self.text_model = AutoModel.from_pretrained("sentence-transformers/all-MiniLM-L6-v2")
                self.text_tokenizer = AutoTokenizer.from_pretrained("sentence-transformers/all-MiniLM-L6-v2")
                self.text_model.to(self.device)
                print("✓ 备用文本编码器加载成功")
            except Exception as e2:
                print(f"⚠ 备用文本编码器加载失败: {e2}")
                self.text_model = None
                self.text_tokenizer = None
    
    def encode_text(self, text: str) -> torch.Tensor:
        """
        编码文本数据
        
        Args:
            text: 输入文本
            
        Returns:
            torch.Tensor: 文本向量表示
        """
        if self.text_model is None:
            raise ValueError("文本编码器未加载成功")
        
        # 使用BGE-M3编码文本
        inputs = self.text_tokenizer(
            text, 
            return_tensors="pt", 
            padding=True, 
            truncation=True, 
            max_length=512
        ).to(self.device)
        
        with torch.no_grad():
            outputs = self.text_model(**inputs)
            # 获取句子级别的嵌入
            if hasattr(outputs, 'last_hidden_state'):
                embeddings = outputs.last_hidden_state[:, 0, :]  # [CLS] token
            else:
                embeddings = outputs.pooler_output
        
        return embeddings.cpu()
    
    def encode_image(self, image_path: str) -> torch.Tensor:
        """
        编码图像数据
        
        Args:
            image_path: 图像文件路径
            
        Returns:
            torch.Tensor: 图像向量表示
        """
        if self.clip_model is None:
            raise ValueError("图像编码器未加载成功")
        
        # 加载并预处理图像
        image = Image.open(image_path).convert("RGB")
        inputs = self.clip_processor(images=image, return_tensors="pt").to(self.device)
        
        with torch.no_grad():
            image_features = self.clip_model.get_image_features(**inputs)
        
        return image_features.cpu()
    
    def encode_table(self, table_path: str) -> torch.Tensor:
        """
        编码表格数据
        
        Args:
            table_path: 表格文件路径 (支持.csv, .xlsx)
            
        Returns:
            torch.Tensor: 表格向量表示
        """
        # 读取表格数据
        ext = os.path.splitext(table_path)[1].lower()
        
        if ext == '.csv':
            df = pd.read_csv(table_path)
        elif ext in ['.xlsx', '.xls']:
            df = pd.read_excel(table_path)
        else:
            raise ValueError(f"不支持的表格格式: {ext}")
        
        # 将表格转换为结构化文本描述
        table_description = self._table_to_text(df)
        
        # 使用文本编码器编码表格描述
        return self.encode_text(table_description)
    
    def _table_to_text(self, df: pd.DataFrame) -> str:
        """将DataFrame转换为结构化文本描述"""
        # 基本表格信息
        text = f"表格包含 {len(df)} 行 {len(df.columns)} 列数据。\n"
        text += f"列名: {', '.join(df.columns.tolist())}\n\n"
        
        # 添加前几行作为样本
        text += "数据样本:\n"
        text += df.head(3).to_markdown(index=False)
        
        # 添加统计信息
        text += "\n\n统计信息:\n"
        numeric_cols = df.select_dtypes(include=['number']).columns
        if len(numeric_cols) > 0:
            text += df[numeric_cols].describe().to_markdown()
        
        return text
    
    def encode_pdf(self, pdf_path: str) -> torch.Tensor:
        """
        编码PDF文档
        
        Args:
            pdf_path: PDF文件路径
            
        Returns:
            torch.Tensor: PDF文档向量表示
        """
        # 提取PDF文本内容
        text = ""
        with pdfplumber.open(pdf_path) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
        
        # 使用文本编码器编码PDF内容
        return self.encode_text(text)
    
    def encode_file(self, file_path: str) -> Dict[str, Any]:
        """
        根据文件类型自动选择编码方式
        
        Args:
            file_path: 文件路径
            
        Returns:
            Dict: 包含编码结果和元数据的字典
        """
        ext = os.path.splitext(file_path)[1].lower()
        
        result = {
            'file_path': file_path,
            'file_type': ext,
            'embedding': None,
            'success': False,
            'error': None
        }
        
        try:
            if ext in ['.txt', '.md', '.json']:
                # 文本文件
                with open(file_path, 'r', encoding='utf-8') as f:
                    text = f.read()
                embedding = self.encode_text(text)
                result['embedding'] = embedding
                result['success'] = True
                
            elif ext in ['.png', '.jpg', '.jpeg', '.bmp', '.gif']:
                # 图像文件
                embedding = self.encode_image(file_path)
                result['embedding'] = embedding
                result['success'] = True
                
            elif ext in ['.csv', '.xlsx', '.xls']:
                # 表格文件
                embedding = self.encode_table(file_path)
                result['embedding'] = embedding
                result['success'] = True
                
            elif ext == '.pdf':
                # PDF文档
                embedding = self.encode_pdf(file_path)
                result['embedding'] = embedding
                result['success'] = True
                
            elif ext in ['.docx', '.doc']:
                # Word文档 - 使用文本编码器
                try:
                    import docx
                    doc = docx.Document(file_path)
                    text = ""
                    for paragraph in doc.paragraphs:
                        text += paragraph.text + "\n"
                    
                    # 添加表格内容
                    for table in doc.tables:
                        for row in table.rows:
                            for cell in row.cells:
                                text += cell.text + "\t"
                            text += "\n"
                        text += "\n"
                    
                    embedding = self.encode_text(text)
                    result['embedding'] = embedding
                    result['success'] = True
                except ImportError:
                    result['error'] = "请安装: pip install python-docx"
                except Exception as e:
                    result['error'] = f"Word文档编码失败: {e}"
                
            else:
                result['error'] = f"不支持的文件格式: {ext}"
                
        except Exception as e:
            result['error'] = str(e)
        
        return result
    
    def batch_encode(self, file_paths: List[str]) -> List[Dict[str, Any]]:
        """批量编码多个文件"""
        results = []
        for file_path in file_paths:
            if os.path.exists(file_path):
                result = self.encode_file(file_path)
                results.append(result)
            else:
                results.append({
                    'file_path': file_path,
                    'success': False,
                    'error': '文件不存在'
                })
        
        return results


def create_encoder() -> MultimodalEncoder:
    """创建多模态编码器实例"""
    return MultimodalEncoder()


# 使用示例
if __name__ == "__main__":
    # 测试编码器
    encoder = MultimodalEncoder()
    
    # 测试不同文件类型
    test_files = [
        "test.txt",      # 文本文件
        "test.jpg",      # 图像文件  
        "test.csv",      # 表格文件
        "test.pdf"       # PDF文档
    ]
    
    for file in test_files:
        if os.path.exists(file):
            result = encoder.encode_file(file)
            if result['success']:
                print(f"{file}: 编码成功, 向量维度: {result['embedding'].shape}")
            else:
                print(f"{file}: 编码失败 - {result['error']}")