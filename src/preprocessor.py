import os
import json
import pdfplumber
import pytesseract
import pandas as pd
from typing import Dict, Any, Optional
from PIL import Image
from .multimodal_encoder import MultimodalEncoder


class DataPreprocessor:
    """增强版数据预处理器，支持多模态编码"""
    
    def __init__(self, use_multimodal_encoding: bool = True):
        """
        初始化预处理器
        
        Args:
            use_multimodal_encoding: 是否启用多模态编码
        """
        self.use_multimodal_encoding = use_multimodal_encoding
        self.multimodal_encoder = None
        
        if use_multimodal_encoding:
            try:
                self.multimodal_encoder = MultimodalEncoder()
                print("✓ 多模态编码器初始化成功")
            except Exception as e:
                print(f"⚠ 多模态编码器初始化失败: {e}")
                self.use_multimodal_encoding = False
    
    def process_file(self, file_path: str) -> Dict[str, Any]:
        """
        处理文件，返回包含文本内容和向量编码的结果
        
        Args:
            file_path: 文件路径
            
        Returns:
            Dict: 包含文本内容和向量编码的字典
        """
        ext = os.path.splitext(file_path)[1].lower()
        
        result = {
            'file_path': file_path,
            'file_type': ext,
            'text_content': '',
            'embedding': None,
            'embedding_dim': 0,
            'success': True,
            'error': None
        }
        
        try:
            # 提取文本内容
            if ext == '.txt':
                with open(file_path, 'r', encoding='utf-8') as f:
                    result['text_content'] = f.read()
            
            elif ext == '.pdf':
                text = ""
                with pdfplumber.open(file_path) as pdf:
                    for page in pdf.pages:
                        page_text = page.extract_text()
                        if page_text:
                            text += page_text + "\n"
                result['text_content'] = text
            
            elif ext in ['.png', '.jpg', '.jpeg', '.bmp']:
                # OCR提取图像中的文本
                try:
                    ocr_text = pytesseract.image_to_string(Image.open(file_path), lang='chi_sim+eng')
                    result['text_content'] = ocr_text
                except Exception as e:
                    result['text_content'] = "[图像文件，OCR识别失败]"
                    result['error'] = f"OCR识别失败: {e}"
            
            elif ext in ['.csv', '.xlsx', '.xls']:
                # 表格文件转换为文本描述
                try:
                    if ext == '.csv':
                        df = pd.read_csv(file_path)
                    else:
                        df = pd.read_excel(file_path)
                    
                    # 创建表格的文本描述
                    table_text = f"表格文件: {os.path.basename(file_path)}\n"
                    table_text += f"数据维度: {len(df)} 行 × {len(df.columns)} 列\n"
                    table_text += f"列名: {', '.join(df.columns.tolist())}\n\n"
                    table_text += "数据预览:\n"
                    table_text += df.head(5).to_markdown(index=False)
                    
                    result['text_content'] = table_text
                except Exception as e:
                    result['text_content'] = "[表格文件，读取失败]"
                    result['error'] = f"表格读取失败: {e}"
            
            elif ext in ['.docx', '.doc']:
                # Word文档支持
                try:
                    # 尝试使用python-docx库
                    import docx
                    doc = docx.Document(file_path)
                    text = ""
                    for paragraph in doc.paragraphs:
                        text += paragraph.text + "\n"
                    
                    # 添加表格内容
                    for table in doc.tables:
                        for row in table.rows:
                            for cell in row.cells:
                                text += cell.text + "\t"
                            text += "\n"
                        text += "\n"
                    
                    result['text_content'] = text
                except ImportError:
                    result['text_content'] = "[Word文档，需要安装python-docx库]"
                    result['error'] = "请安装: pip install python-docx"
                except Exception as e:
                    result['text_content'] = "[Word文档，读取失败]"
                    result['error'] = f"Word文档读取失败: {e}"
            
            else:
                result['success'] = False
                result['error'] = f"不支持的文件格式: {ext}"
                return result
            
            # 如果启用了多模态编码，生成向量嵌入
            if self.use_multimodal_encoding and self.multimodal_encoder:
                try:
                    encoding_result = self.multimodal_encoder.encode_file(file_path)
                    if encoding_result['success']:
                        result['embedding'] = encoding_result['embedding']
                        if result['embedding'] is not None:
                            result['embedding_dim'] = result['embedding'].shape[-1]
                except Exception as e:
                    result['error'] = f"向量编码失败: {e}"
        
        except Exception as e:
            result['success'] = False
            result['error'] = str(e)
        
        return result
    
    def process_files(self, file_paths: list) -> list:
        """批量处理多个文件"""
        results = []
        for file_path in file_paths:
            if os.path.exists(file_path):
                result = self.process_file(file_path)
                results.append(result)
            else:
                results.append({
                    'file_path': file_path,
                    'success': False,
                    'error': '文件不存在'
                })
        
        return results
    
    def get_text_only(self, file_path: str) -> str:
        """仅获取文本内容（兼容旧版本）"""
        result = self.process_file(file_path)
        return result['text_content'] if result['success'] else ""


# 兼容旧版本的静态方法
class LegacyPreprocessor:
    """兼容旧版本的预处理器"""
    
    @staticmethod
    def process_file(file_path: str) -> str:
        """根据文件扩展名选择不同的解析方式，转换为统一文本"""
        ext = os.path.splitext(file_path)[1].lower()
        if ext == '.txt':
            with open(file_path, 'r', encoding='utf-8') as f:
                return f.read()
        elif ext == '.pdf':
            text = ""
            with pdfplumber.open(file_path) as pdf:
                for page in pdf.pages:
                    text += page.extract_text() + "\n"
            return text
        elif ext in ['.png', '.jpg', '.jpeg']:
            # 需要提前安装 tesseract-ocr
            return pytesseract.image_to_string(Image.open(file_path), lang='chi_sim')
        else:
            raise ValueError(f"暂不支持的文件格式: {ext}")